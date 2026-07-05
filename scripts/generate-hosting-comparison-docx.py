"""Generate ProVizient hosting comparison Word document for client review."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, highlight_row: int | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
        set_cell_shading(hdr_cells[i], "0066FF")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
        if highlight_row is not None and r_idx == highlight_row:
            for cell in row_cells:
                set_cell_shading(cell, "E8F4FF")
    doc.add_paragraph()
    return table


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("ProVizient Website Hosting", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Provider Comparison & Recommendation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared for: Client Decision Review\n").italic = True
    meta.add_run("Company: ProVizient | Website: provizient.com\n").italic = True
    meta.add_run("Date: June 2026 | Document: Hosting Options Analysis").italic = True
    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "ProVizient requires production hosting for a modern marketing website with contact forms, "
        "custom domain (provizient.com), business email (info@ProVizient.com), and SSL security. "
        "This document compares Azure, Hostinger, and Hosting.com across cost, security, and technical fit."
    )
    rec = doc.add_paragraph()
    rec.add_run("Recommended option: Microsoft Azure (Static Web Apps Standard + Microsoft 365 Basic)\n").bold = True
    rec.add_run(
        "Estimated cost: ~$16/month (~$192/year) — predictable pricing, enterprise security, "
        "and native support for the existing website codebase with no redevelopment required."
    )
    doc.add_paragraph()

    doc.add_heading("1. Important Technical Note", level=1)
    add_table(
        doc,
        ["Provider Type", "Fits Current ProVizient Site?"],
        [
            [
                "Azure Static Web Apps",
                "BEST FIT — Site is already built for Azure (static export + Azure Functions for forms).",
            ],
            [
                "Hostinger / Hosting.com",
                "PARTIAL FIT — Static pages only; contact forms require rewrite (PHP or third-party email).",
            ],
        ],
    )
    doc.add_paragraph(
        "Conclusion: The lowest promotional price is not always the best value if code changes, "
        "form rework, and renewal price increases are required."
    )

    doc.add_heading("2. What Is Included in This Comparison", level=1)
    doc.add_paragraph("Included: Web hosting, custom domain, business email, SSL certificate.")
    doc.add_paragraph("Excluded from base comparison: Optional 20 GB database (see Section 6).")
    doc.add_paragraph("Current demo hosting: GitHub Pages (free, for preview only — not recommended for production).")

    doc.add_heading("3. Year 1 Cost (Intro / Promotional Pricing)", level=1)
    add_table(
        doc,
        ["Provider", "Plan", "Hosting/mo", "Hosting/yr", "Domain/yr", "Email/yr", "Total/mo", "Total/yr"],
        [
            ["Azure", "SWA Standard + M365 Basic", "$9", "$108", "~$12", "~$72", "~$16", "~$192"],
            ["Hostinger", "Business (long-term promo)*", "~$3–4", "~$36–48", "$0 (1st yr)", "$0 (1st yr)", "~$3–4", "~$36–48"],
            ["Hosting.com", "Shared Plus (intro)", "~$4", "~$48", "$0–10", "varies", "~$4–5", "~$48–60"],
        ],
        highlight_row=0,
    )
    doc.add_paragraph(
        "* Hostinger intro pricing typically requires 1–4 years paid upfront "
        "(e.g., ~$191 for 48 months ≈ $3.99/month effective)."
    )

    doc.add_heading("4. Year 2+ Cost (Renewal / Steady Pricing)", level=1)
    add_table(
        doc,
        ["Provider", "Hosting/mo", "Hosting/yr", "Domain/yr", "Email/yr", "Total/mo", "Total/yr"],
        [
            ["Azure", "$9", "$108", "~$12", "~$72", "~$16", "~$192"],
            ["Hostinger", "~$17", "~$204", "~$20", "~$19+", "~$20–22", "~$243+"],
            ["Hosting.com", "~$11–15", "~$132–180", "~$15", "~$12–60", "~$15–18", "~$180–220"],
        ],
        highlight_row=0,
    )

    doc.add_heading("5. Full Stack Comparison", level=1)
    add_table(
        doc,
        ["Criteria", "Azure", "Hostinger Business", "Hosting.com Shared"],
        [
            ["Year 1 total (no database)", "~$192", "~$36–48", "~$48–60"],
            ["Year 2+ total (no database)", "~$192 (stable)", "~$243+ (renewal jump)", "~$180–220"],
            ["Price predictable?", "High", "Low", "Medium"],
            ["Security rating", "Enterprise (5/5)", "Good SMB (3/5)", "Good SMB (3/5)"],
            ["Free SSL", "Yes", "Yes", "Yes"],
            ["DDoS / Enterprise SLA", "Yes (Standard plan)", "Basic", "Basic"],
            ["Fits Next.js + API forms", "Yes (native)", "No (rewrite needed)", "No (rewrite needed)"],
            ["Custom domain", "Yes", "Yes", "Yes"],
            ["Business email", "M365 (~$6/mo)", "Included (promo)", "Add-on / limited"],
            ["Best for current codebase", "Yes", "Needs changes", "Needs changes"],
        ],
        highlight_row=None,
    )

    doc.add_heading("6. Optional: Database (~20 GB Storage)", level=1)
    add_table(
        doc,
        ["Provider", "Database Service", "Extra/mo", "Extra/yr", "Notes"],
        [
            ["Azure", "PostgreSQL Flexible Server (32 GB min)", "~$18–30", "~$216–360", "Managed, enterprise-grade"],
            ["Hostinger", "MySQL (shared, included)", "$0", "$0", "Included in Business plan"],
            ["Hosting.com", "MySQL/MariaDB (shared)", "$0", "$0", "Included in plan"],
        ],
    )
    doc.add_paragraph(
        "Note: Azure PostgreSQL minimum storage is 32 GB (not 20 GB). Shared MySQL on budget hosts "
        "is suitable for simple applications but is not equivalent to managed PostgreSQL for enterprise use."
    )

    doc.add_heading("7. Three-Year Total Cost Estimate", level=1)
    add_table(
        doc,
        ["Provider", "Approximate 3-Year Total", "Notes"],
        [
            ["Azure (site + domain + email, no DB)", "~$576", "Stable annual cost"],
            ["Hostinger (intro + 2 renewal years)", "~$534", "Year 1 cheap, renewals increase sharply"],
            ["Hosting.com", "~$450", "Mid-range after promotions end"],
            ["Azure + PostgreSQL database", "~$1,200+", "Full enterprise stack"],
        ],
        highlight_row=0,
    )

    doc.add_heading("8. Security Comparison", level=1)
    doc.add_paragraph("1. Azure — Enterprise platform, SLA, identity integration, managed services, DDoS protection.")
    doc.add_paragraph("2. Hosting.com — Solid shared hosting with malware protection.")
    doc.add_paragraph("3. Hostinger — Good for small sites; shared server environment (multi-tenant).")
    doc.add_paragraph(
        "For a business consulting brand serving enterprise clients, Azure provides the strongest "
        "security posture and professional credibility."
    )

    doc.add_heading("9. Decision Matrix — Quick Reference", level=1)
    add_table(
        doc,
        ["If you want…", "Best choice"],
        [
            ["Cheapest demo / proof of concept", "Keep GitHub Pages (free)"],
            ["Cheapest paid hosting (static only)", "Hostinger Business (Year 1 promo)"],
            ["Production site + forms + domain + email", "Azure (~$16/mo)"],
            ["Maximum security + future database/admin", "Azure (+ DB later ~$20–30/mo)"],
            ["Lowest 3-year cost with no code changes", "Azure"],
        ],
        highlight_row=3,
    )

    doc.add_heading("10. Official Recommendation", level=1)
    p = doc.add_paragraph()
    p.add_run("RECOMMENDED: Microsoft Azure\n\n").bold = True
    p.add_run(
        "We recommend Azure Static Web Apps (Standard plan) combined with Microsoft 365 Business Basic "
        "for the info@ProVizient.com mailbox.\n\n"
        "Reasons:\n"
        "• The website is already engineered for Azure — zero redevelopment cost\n"
        "• Contact and consultation forms work natively via Azure Functions\n"
        "• Predictable ~$192/year with no surprise renewal price jumps\n"
        "• Enterprise-grade security, SSL, and SLA suitable for client-facing business\n"
        "• Easy path to add PostgreSQL database later if lead storage or admin panel is needed\n"
        "• Professional business email with info@ProVizient.com on your own domain\n\n"
        "Alternative if budget is the only priority in Year 1: Hostinger Business (~$36–48 first year) "
        "accepting that forms must be rebuilt and renewal costs rise to ~$243+/year.\n\n"
        "Not recommended for production: Remaining on GitHub Pages (demo only, no business email, forms in demo mode)."
    )

    doc.add_heading("11. Next Steps (Upon Client Approval)", level=1)
    steps = [
        "Purchase domain provizient.com via Azure App Service Domains (~$12/year).",
        "Create Azure Static Web Apps (Standard) and connect GitHub repository.",
        "Configure custom domain and free SSL certificate.",
        "Set up Microsoft 365 Business Basic for info@ProVizient.com.",
        "Configure Azure Communication Services or Resend for form email delivery.",
        "Disable GitHub Pages demo or keep as staging preview only.",
        "(Optional Phase 2) Add Azure PostgreSQL if database and admin panel are required.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    doc.add_paragraph()
    footer = doc.add_paragraph("— End of Document —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True

    out_path = r"d:\provizient\docs\ProVizient-Hosting-Comparison.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
