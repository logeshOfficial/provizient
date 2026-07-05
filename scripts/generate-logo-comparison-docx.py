"""Generate ProVizient logo comparison Word document for client finalization."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"d:\provizient")
LOGO_DIR = ROOT / "public" / "logo-options"
PREVIEW_DIR = ROOT / "docs" / "logo-previews"
OUT_PATH = ROOT / "docs" / "ProVizient-Logo-Comparison.docx"
OUT_PATH_FALLBACK = ROOT / "docs" / "ProVizient-Logo-Comparison-Updated.docx"

OPTIONS = [
    {
        "num": 1,
        "name": "Neural Network",
        "file": "provizient-logo-option-1-neural-network.png",
        "preview": "option-1-neural-network-website.png",
        "concept": (
            "A multi-color neural node grid (red, green, blue, yellow) connected by "
            "thin lines, paired with the ProVizient wordmark and full tagline."
        ),
        "strengths": (
            "Professional and credible; closest to the original brand direction; "
            "reads clearly at small sizes in website headers and documents."
        ),
        "best_for": "Consulting, enterprise clients, and a balanced AI + software brand.",
        "tone": "Professional / Corporate",
    },
    {
        "num": 2,
        "name": "Circuit P",
        "file": "provizient-logo-option-2-circuit-p.png",
        "preview": "option-2-circuit-p-website.png",
        "concept": (
            "A stylized letter \"P\" built from circuit-board traces and nodes, "
            "with the ProVizient wordmark and tagline."
        ),
        "strengths": (
            "Strong technology identity; memorable monogram; clean and modern without "
            "being playful."
        ),
        "best_for": "Software development emphasis and a sharper tech-company look.",
        "tone": "Modern / Tech-forward",
    },
    {
        "num": 3,
        "name": "Robot Hex",
        "file": "provizient-logo-option-3-robot-hex.png",
        "preview": "option-3-robot-hex-website.png",
        "concept": (
            "A friendly robot head inside a hexagonal badge, paired with the "
            "ProVizient wordmark and tagline."
        ),
        "strengths": (
            "Distinct AI and agentic-AI personality; aligns with the website's "
            "innovation hub and training themes."
        ),
        "best_for": "AI training, agentic AI products, and a more approachable brand.",
        "tone": "Innovative / Friendly",
    },
    {
        "num": 4,
        "name": "Growth + AI",
        "file": "provizient-logo-option-4-growth-ai.png",
        "preview": "option-4-growth-ai-website.png",
        "concept": (
            "Neural pathways with an upward growth arrow, symbolizing career "
            "advancement through AI learning."
        ),
        "strengths": (
            "Emphasizes training and career outcomes; positive and aspirational; "
            "good fit for education-focused messaging."
        ),
        "best_for": "Training programs, career development, and workforce upskilling.",
        "tone": "Aspirational / Education",
    },
    {
        "num": 5,
        "name": "Data Mark",
        "file": "provizient-logo-option-5-data-mark.png",
        "preview": "option-5-data-mark-website.png",
        "concept": (
            "A rounded-square icon with colorful data bars and analytics motif, "
            "with the ProVizient wordmark and tagline."
        ),
        "strengths": (
            "Enterprise and analytics feel; structured and confident; works well "
            "for data-driven consulting positioning."
        ),
        "best_for": "Enterprise analytics, business intelligence, and consulting services.",
        "tone": "Enterprise / Analytics",
    },
]


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, highlight_row: int | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
        set_cell_shading(hdr_cells[i], "0066FF")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
        if highlight_row is not None and r_idx == highlight_row:
            for cell in row_cells:
                set_cell_shading(cell, "E8F4FF")
    doc.add_paragraph()
    return table


def add_centered_picture(doc, image_path: Path, width_inches: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def add_option_section(doc, option: dict) -> None:
    doc.add_heading(f"Option {option['num']} — {option['name']}", level=2)

    p = doc.add_paragraph()
    p.add_run("Logo artwork").bold = True

    image_path = LOGO_DIR / option["file"]
    if image_path.exists():
        add_centered_picture(doc, image_path, 5.8)
    else:
        doc.add_paragraph(f"[Logo image not found: {option['file']}]")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Website preview (header + homepage)").bold = True

    preview_path = PREVIEW_DIR / option["preview"]
    if preview_path.exists():
        add_centered_picture(doc, preview_path, 6.2)
        cap = doc.add_paragraph(
            "Screenshot showing this logo applied to the ProVizient website header and hero section."
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph(f"[Website preview not found: {option['preview']}]")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Concept: ").bold = True
    p.add_run(option["concept"])

    p = doc.add_paragraph()
    p.add_run("Strengths: ").bold = True
    p.add_run(option["strengths"])

    p = doc.add_paragraph()
    p.add_run("Best suited for: ").bold = True
    p.add_run(option["best_for"])

    p = doc.add_paragraph()
    p.add_run("Brand tone: ").bold = True
    p.add_run(option["tone"])

    doc.add_paragraph()


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("ProVizient Logo Concepts", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Client Review & Final Selection")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared for: Client Decision Review\n").italic = True
    meta.add_run("Company: ProVizient | Website: provizient.com\n").italic = True
    meta.add_run("Date: June 2026 | Document: Logo Options Comparison").italic = True
    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This document presents five logo concepts developed for ProVizient's new website "
        "and brand identity. Each option includes the company name, tagline "
        "(AI Training | Software Development | Consulting), and a distinct icon style "
        "reflecting a different brand emphasis."
    )
    doc.add_paragraph(
        "All concepts have been applied to the live website preview for side-by-side "
        "evaluation in the header and footer. Please review each option below and "
        "indicate your preferred choice in the Client Selection section at the end of "
        "this document."
    )
    doc.add_paragraph(
        "Live website preview (demo): https://logeshOfficial.github.io/provizient/"
    )
    doc.add_paragraph()

    doc.add_heading("Brand Context", level=1)
    add_table(
        doc,
        ["Element", "Details"],
        [
            ["Company name", "ProVizient"],
            ["Tagline", "Empowering Careers. Elevating Futures."],
            ["Service lines", "AI Training | Software Development | Consulting"],
            ["Business email", "info@ProVizient.com"],
            ["Location", "12205 Settlers Drive, Frisco, TX 75035"],
            ["Logo usage", "Website header, footer, business documents, email signature, social profiles"],
        ],
    )

    doc.add_heading("Quick Comparison", level=1)
    add_table(
        doc,
        ["Option", "Name", "Icon Style", "Brand Tone", "Best For"],
        [
            [
                "1",
                "Neural Network",
                "Connected AI nodes",
                "Professional / Corporate",
                "Consulting & enterprise credibility",
            ],
            [
                "2",
                "Circuit P",
                "Circuit-board monogram",
                "Modern / Tech-forward",
                "Software development focus",
            ],
            [
                "3",
                "Robot Hex",
                "Hex badge + robot",
                "Innovative / Friendly",
                "AI training & agentic AI",
            ],
            [
                "4",
                "Growth + AI",
                "Neural paths + arrow",
                "Aspirational / Education",
                "Career growth & training",
            ],
            [
                "5",
                "Data Mark",
                "Analytics bars",
                "Enterprise / Analytics",
                "Data-driven consulting",
            ],
        ],
        highlight_row=0,
    )

    doc.add_heading("Logo Concepts (Detailed)", level=1)
    doc.add_paragraph(
        "Each option includes the standalone logo artwork and a website screenshot showing "
        "how it appears in the site header and homepage hero section."
    )
    for option in OPTIONS:
        add_option_section(doc, option)

    doc.add_heading("Developer Recommendation", level=1)
    p = doc.add_paragraph()
    p.add_run("Recommended for client consideration: Option 1 — Neural Network\n\n").bold = True
    p.add_run(
        "Option 1 offers the strongest balance of professionalism and AI identity for "
        "a consulting and training firm serving business clients. It aligns with the "
        "existing brand direction, remains legible at small sizes in the website header, "
        "and presents a credible image for proposals, email, and print materials.\n\n"
        "Option 3 (Robot Hex) is a strong alternative if ProVizient wants a more "
        "approachable, product-led AI brand. Options 2, 4, and 5 are viable depending "
        "on whether the priority is software engineering (2), training outcomes (4), "
        "or enterprise analytics (5)."
    )

    doc.add_heading("Client Selection", level=1)
    doc.add_paragraph(
        "Please mark your preferred logo option and return this section (or reply by email) "
        "to confirm your final choice."
    )
    add_table(
        doc,
        ["Option", "Logo Name", "Select (✓)"],
        [[str(o["num"]), o["name"], ""] for o in OPTIONS],
    )

    doc.add_paragraph()
    doc.add_paragraph("Selected option number: __________")
    doc.add_paragraph("Approved by (name): _________________________________")
    doc.add_paragraph("Title / role: ______________________________________")
    doc.add_paragraph("Date: ______________________________________________")
    doc.add_paragraph()
    doc.add_paragraph("Optional notes or revision requests:")
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)

    doc.add_heading("Next Steps (Upon Client Approval)", level=1)
    steps = [
        "Apply the selected logo as the production site logo (header, footer, and favicon).",
        "Export final logo files: full-color PNG, transparent background where applicable, and high-resolution PDF.",
        "Update business email signature and any existing marketing collateral.",
        "Deploy the finalized logo to the production website (provizient.com).",
        "Archive alternate concepts for potential future brand extensions.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    doc.add_paragraph()
    footer = doc.add_paragraph("— End of Document —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUT_PATH))
        print(f"Saved: {OUT_PATH}")
    except PermissionError:
        doc.save(str(OUT_PATH_FALLBACK))
        print(f"Original file is open — saved instead to: {OUT_PATH_FALLBACK}")


if __name__ == "__main__":
    main()
